import streamlit as st
import pdfplumber
import pandas as pd
from io import BytesIO
from datetime import datetime
from docx import Document

# دوال تحليل بيانات الحضور والانصراف
def is_morning_shift(arrival_time):
    return 9 <= arrival_time.hour < 14  # وردية صباحية بين 9:00 صباحًا و 1:59 مساءً

def is_evening_shift(arrival_time):
    return arrival_time.hour >= 14  # وردية مسائية بعد الساعة 2:00 مساءً

def is_double_shift(arrival_time, departure_time):
    return arrival_time.hour < 14 and departure_time.hour >= 22  # وردية مزدوجة، الحضور قبل 2:00 والخروج بعد 10:00 مساءً

def is_single_punch_shift(punch_times):
    return len(punch_times) == 1  # وردية بصمة واحدة إذا كانت بصمة واحدة فقط

def count_shifts(shift_type):
    if shift_type in ['مسائية', 'صباحية', 'بصمة واحدة']:
        return 1
    elif shift_type == 'مزدوجة':
        return 2  # حساب عدد الورديات

def calculate_overtime(total_shifts):
    return total_shifts - 26 if total_shifts > 26 else 0  # حساب الورديات الشهرية

def count_delays(punch_times):
    # التأخيرات بين 3:10 مساءً و4:59 مساءً فقط
    delays = [p for p in punch_times if (p.hour == 15 and p.minute > 10) or (p.hour == 16)]
    return len(delays), delays

# دالة لتحليل بيانات الحضور والانصراف من ملف PDF
def load_pdf_data(file):
    data = []
    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            table = page.extract_table()
            if table:
                data.extend(table[1:])  # تجاهل رؤوس الأعمدة
    return data

def format_time_12h(time):
    return time.strftime("%I:%M %p")  # تحويل الوقت إلى نظام 12 ساعة

def process_attendance_data(data):
    # تحويل البيانات إلى DataFrame
    df = pd.DataFrame(data, columns=["رقم البصمه", "الإسم", "داعيملا"])
    df['التاريخ_الوقت'] = pd.to_datetime(df['داعيملا'], errors='coerce')
    df.dropna(subset=['التاريخ_الوقت'], inplace=True)  # إزالة الصفوف غير الصالحة

    summary = []
    for name, group in df.groupby("الإسم"):
        group = group.sort_values('التاريخ_الوقت')  # فرز حسب التاريخ

        shifts_count, delays_count, overtime_count = 0, 0, 0
        shift_types, delay_details = set(), []
        overtime_details = []  # لتخزين تفاصيل الدوام الإضافي

        for day, punches in group.groupby(group['التاريخ_الوقت'].dt.date):
            punch_times = punches['التاريخ_الوقت'].tolist()
            arrival, departure = min(punch_times), max(punch_times)
            
            # تحديد نوع الوردية
            if is_single_punch_shift(punch_times):
                shift_type = 'بصمة واحدة'
                shifts_count += count_shifts(shift_type)
            elif is_double_shift(arrival, departure):
                shift_type = 'مزدوجة'
                shifts_count += count_shifts(shift_type)
            elif is_morning_shift(arrival):
                shift_type = 'صباحية'
                shifts_count += count_shifts(shift_type)
            elif is_evening_shift(arrival):
                shift_type = 'مسائية'
                shifts_count += count_shifts(shift_type)
            
            shift_types.add(shift_type)
            
            # حساب التأخيرات
            daily_delays, delay_times_list = count_delays(punch_times)
            delays_count += daily_delays
            for delay_time in delay_times_list:
                delay_details.append([str(day), format_time_12h(delay_time)])

            # تخزين تفاصيل الدوام الإضافي
            if shifts_count > 26:
                overtime_details.append([str(day), format_time_12h(arrival), format_time_12h(departure), shift_type])

        overtime_count = calculate_overtime(shifts_count)

        summary.append({
            "الإسم": name,
            "عدد الورديات": shifts_count,
            "نوع الورديات": ', '.join(shift_types),
            "عدد التأخيرات": delays_count,
            "تفاصيل التأخيرات": delay_details,
            "أيام الغياب": max(0, 26 - shifts_count),
            "الدوام الإضافي": overtime_count,
            "تفاصيل الدوام الإضافي": overtime_details  # إضافة تفاصيل الدوام الإضافي
        })

    return pd.DataFrame(summary)

# واجهة Streamlit
st.title("تحليل بيانات الحضور والانصراف")
st.write("قم بتحميل ملف PDF لتحليل بيانات الحضور والانصراف")

uploaded_file = st.file_uploader("اختر ملف PDF", type="pdf")

if uploaded_file is not None:
    st.write("جاري معالجة الملف...")
    data = load_pdf_data(uploaded_file)
    
    if data:
        summary_df = process_attendance_data(data)
        st.write("النتائج:")
        st.dataframe(summary_df)  # عرض البيانات في واجهة Streamlit
        
        # تحميل التقرير كملف Word
        output = BytesIO()
        doc = Document()
        doc.add_heading("تقرير الحضور والانصراف", 0)
        
        for index, row in summary_df.iterrows():
            paragraph = doc.add_paragraph()
            paragraph.add_run(f"الموظف: {row['الإسم']}\n").bold = True
            paragraph.add_run(f"عدد الورديات: {row['عدد الورديات']}\n")
            paragraph.add_run(f"نوع الورديات: {row['نوع الورديات']}\n")
            paragraph.add_run(f"عدد التأخيرات: {row['عدد التأخيرات']}\n")
            paragraph.add_run(f"\nأيام الغياب: {row['أيام الغياب']}\n")
            paragraph.add_run(f"الدوام الإضافي: {row['الدوام الإضافي']}\n")

        # حفظ التقرير كملف Word وتنزيله
        doc.save(output)
        output.seek(0)
        st.download_button(label="تحميل التقرير كملف Word", data=output, file_name="attendance_report.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    else:
        st.write("لم يتم العثور على بيانات في الملف. يرجى التأكد من صحة التنسيق.")

